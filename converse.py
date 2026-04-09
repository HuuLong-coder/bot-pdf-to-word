import argparse
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document

try:
    from pdf2docx import Converter as Pdf2DocxConverter
except Exception:  # noqa: BLE001
    Pdf2DocxConverter = None


@dataclass
class ConversionResult:
    pages_total: int
    pages_with_text: int
    output_path: Path
    engine_used: str


def _normalize_lines(text: str) -> list[str]:
    """Làm sạch dữ liệu text để ghi ra Word dễ đọc hơn."""
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split())
        if line:
            lines.append(line)
    return lines


def _count_pdf_pages(pdf_file: Path) -> int:
    with pdfplumber.open(str(pdf_file)) as pdf:
        return len(pdf.pages)


def _convert_with_pdf2docx(pdf_file: Path, output_file: Path) -> ConversionResult:
    if Pdf2DocxConverter is None:
        raise RuntimeError(
            "Chưa cài thư viện pdf2docx. Chạy: python -m pip install pdf2docx"
        )

    pages_total = _count_pdf_pages(pdf_file)
    if pages_total == 0:
        raise ValueError("File PDF không có trang nào.")

    converter = Pdf2DocxConverter(str(pdf_file))
    try:
        converter.convert(str(output_file), start=0, end=None)
    finally:
        converter.close()

    return ConversionResult(
        pages_total=pages_total,
        pages_with_text=pages_total,
        output_path=output_file,
        engine_used="pdf2docx",
    )


def _convert_with_text_extraction(
    pdf_file: Path,
    output_file: Path,
    *,
    title: str,
    add_page_headings: bool,
) -> ConversionResult:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    pages_with_text = 0

    with pdfplumber.open(str(pdf_file)) as pdf:
        if not pdf.pages:
            raise ValueError("File PDF không có trang nào.")

        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()

            if add_page_headings:
                doc.add_heading(f"Trang {i}", level=2)

            if text and text.strip():
                pages_with_text += 1
                lines = _normalize_lines(text)
                for line in lines:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph("[Không trích xuất được nội dung ở trang này]")

            if i < len(pdf.pages):
                doc.add_page_break()

    doc.save(str(output_file))
    return ConversionResult(
        pages_total=len(pdf.pages),
        pages_with_text=pages_with_text,
        output_path=output_file,
        engine_used="text-extraction",
    )


def pdf_to_word(
    pdf_path: str | Path,
    docx_path: str | Path,
    *,
    title: str = "Nội dung chuyển từ PDF",
    add_page_headings: bool = True,
    engine: str = "auto",
) -> ConversionResult:
    """
    Chuyển nội dung text từ PDF sang Word (.docx).

    engine:
    - auto: thử giữ bố cục bằng pdf2docx, lỗi thì fallback text extraction.
    - layout: chỉ dùng pdf2docx (giữ layout tốt hơn).
    - text: trích text thuần (nhanh, nhưng mất layout).
    """

    pdf_file = Path(pdf_path)
    if not pdf_file.exists():
        raise FileNotFoundError(f"Không tìm thấy file PDF: {pdf_path}")
    if pdf_file.suffix.lower() != ".pdf":
        raise ValueError(f"File đầu vào không phải PDF: {pdf_path}")

    output_file = Path(docx_path)
    if output_file.suffix.lower() != ".docx":
        raise ValueError(f"File đầu ra phải có đuôi .docx: {docx_path}")
    output_file.parent.mkdir(parents=True, exist_ok=True)

    normalized_engine = engine.lower().strip()
    if normalized_engine not in {"auto", "layout", "text"}:
        raise ValueError("engine phải là một trong các giá trị: auto, layout, text")

    if normalized_engine == "layout":
        return _convert_with_pdf2docx(pdf_file, output_file)

    if normalized_engine == "text":
        return _convert_with_text_extraction(
            pdf_file,
            output_file,
            title=title,
            add_page_headings=add_page_headings,
        )

    try:
        return _convert_with_pdf2docx(pdf_file, output_file)
    except Exception:
        return _convert_with_text_extraction(
            pdf_file,
            output_file,
            title=title,
            add_page_headings=add_page_headings,
        )


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Chuyển PDF sang Word với 2 engine: giữ layout (pdf2docx) hoặc text extraction"
    )
    parser.add_argument("pdf_input", help="Đường dẫn file PDF đầu vào")
    parser.add_argument(
        "-o",
        "--output",
        dest="word_output",
        default=None,
        help="Đường dẫn file Word đầu ra (.docx). Mặc định: cùng tên với file PDF",
    )
    parser.add_argument(
        "--title",
        default="Nội dung chuyển từ PDF",
        help="Tiêu đề chính của tài liệu Word",
    )
    parser.add_argument(
        "--no-page-headings",
        action="store_true",
        help="Tắt heading theo từng trang",
    )
    parser.add_argument(
        "--engine",
        choices=["auto", "layout", "text"],
        default="auto",
        help="Engine convert: auto (mac dinh), layout (dep hon), text (don gian)",
    )
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    parser = _build_arg_parser()
    args = parser.parse_args(argv)

    pdf_input = Path(args.pdf_input)
    word_output = (
        Path(args.word_output)
        if args.word_output
        else pdf_input.with_suffix(".docx")
    )

    try:
        result = pdf_to_word(
            pdf_path=pdf_input,
            docx_path=word_output,
            title=args.title,
            add_page_headings=not args.no_page_headings,
            engine=args.engine,
        )
        print(
            "Chuyển đổi thành công | "
            f"Engine: {result.engine_used} | "
            f"Tổng trang: {result.pages_total} | "
            f"Trang có text: {result.pages_with_text} | "
            f"Output: {result.output_path}"
        )
        return 0
    except Exception as exc:
        print(f"Lỗi: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())